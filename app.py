import json
import os
import uuid
from contextlib import nullcontext
from io import BytesIO
import re
from typing import Any, Dict, Iterable, List, Optional, Tuple

import streamlit as st

try:
    import google.generativeai as genai
except ImportError as exc:  # pragma: no cover - streamlit will surface this to the user
    genai = None
    _import_error = exc
else:
    _import_error = None

try:
    from google.api_core.exceptions import NotFound
except ImportError:  # pragma: no cover - optional dependency
    NotFound = Exception

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    Document = None
    Inches = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover - optional dependency
    Image = None

try:
    from streamlit.runtime.secrets import StreamlitSecretNotFoundError
except ImportError:  # pragma: no cover
    StreamlitSecretNotFoundError = Exception  # type: ignore

BASE_PROMPT = """あなたは「めざましメディア」の編集記者です。\n以下の入力データをもとに、めざましメディア風のリリース記事を作成してください。\n\n## 出力ルール\n- リード文 → 小見出し(h2) → 本文詳細 → コメント/反響 → 公式情報ボックス → まとめ\n- 句読点はシンプルに。「！」や「…」も自然な範囲で使用\n- 写真キャプションは「◯◯する△△」の形式で具体的に\n- 読者の感情を引きつける「かわいい」「注目」「大反響」などのワードを適度に盛り込む\n- 最後は「ぜひチェックしてみてください」「お見逃しなく！」などで締める\n\n---\n\n## 入力データ\n- 【タイトル】：\n- 【主役（人物/キャラクター/ブランドなど）】：\n- 【発売日/公開日/開始日】：\n- 【開催場所/販売場所】：\n- 【イベント/商品/作品の特徴】：\n- 【コメントやSNS反応】：\n- 【写真リスト（キャプション用）】：\n- 【公式情報（価格・日程・注意事項など）】：\n\n---\n\n## 記事構造（生成する文章の型）\n\n### 1. リード文（冒頭パラグラフ）\n- 誰が・何を・いつ行うかを端的に\n- 必要に応じてSNSや話題性を一文追加\n\n### 2. 小見出し（h2）\n- 注目ポイントをキャッチーに表現\n  （例：「◯◯あふれる先行カット公開」「かわいすぎる◯◯が新登場！」）\n\n### 3. 本文詳細\n- イベントや商品の背景、ラインナップ、見どころを小分けに説明\n- 写真とキャプションを数点挿入（文章の中で「◯◯する△△」の形で）\n\n### 4. コメント・反響\n- 本人や関係者のコメントを引用\n- SNSの声（例：「かわいすぎる！」「絶対欲しい」など）を紹介\n\n### 5. 公式情報（ボックス形式）\n- 「■販売期間」「■価格」「■場所」などを箇条書きで明記\n\n### 6. まとめ\n- 「ぜひチェックしてみてください」「お見逃しなく！」などで読者を誘導\n\n---\n\n## 出力例フォーマット（イメージ）\n\n<h2>◯◯◯◯</h2>\n\n<p>リード文…</p>\n\n<h2>小見出し</h2>\n<p>詳細説明…</p>\n<figcaption>キャプション例：笑顔を見せる◯◯</figcaption>\n\n<h2>コメント・反響</h2>\n<p>◯◯さんのコメント「……」</p>\n<p>SNSでは「……」といった声も。</p>\n\n<div class=\"mezamashi-box\">\n<p>■発売日：◯月◯日<br>\n■価格：◯円<br>\n■場所：◯◯</p>\n</div>\n\n<p>ぜひチェックしてみてください！</p>"""

OUTPUT_FORMAT_INSTRUCTIONS = """# 出力フォーマット\n必ず JSON 形式のみで回答してください。コードブロックや追加の説明文は一切付けないでください。\n構造は次の通りです。\n{\n  \"headlines\": [\"見出し案1\", ..., \"見出し案10\"],\n  \"subheadlines\": [\"小見出し案1\", ..., \"小見出し案10\"],\n  \"article\": \"本文全体（純粋なプレーンテキスト。HTMLタグやMarkdown記法を含めない）\"\n}\n- \"headlines\" と \"subheadlines\" の配列は必ず10個の要素を含めてください。\n- 文字列内の改行は \n で表現し、ダブルクォートはエスケープしてください。\n- HTMLタグやMarkdown記法を含めず、プレーンテキストのみで記述してください。\n- JSON 以外の文字・コメントは出力しないでください。"""

IMAGE_SUFFIXES = {"png", "jpg", "jpeg", "webp", "bmp", "gif"}
HTML_TAG_RE = re.compile(r"<[^>]+>")


MODEL_OPTIONS = {
    "Gemini 2.5 Pro": [
        "models/gemini-2.5-pro",
        "gemini-2.5-pro",
        "models/gemini-2.5-pro-preview-06-05",
        "gemini-2.5-pro-preview-06-05",
        "models/gemini-2.5-pro-preview-05-06",
        "gemini-2.5-pro-preview-05-06",
        "models/gemini-2.5-pro-preview-03-25",
        "gemini-2.5-pro-preview-03-25",
    ],
    "Gemini 2.5 Flash": [
        "models/gemini-2.5-flash",
        "gemini-2.5-flash",
        "models/gemini-2.5-flash-preview-09-2025",
        "gemini-2.5-flash-preview-09-2025",
        "models/gemini-2.5-flash-preview-05-20",
        "gemini-2.5-flash-preview-05-20",
        "models/gemini-2.5-flash-preview-03-25",
        "gemini-2.5-flash-preview-03-25",
    ],
    "Gemini 2.5 Flash Lite": [
        "models/gemini-2.5-flash-lite",
        "gemini-2.5-flash-lite",
        "models/gemini-2.5-flash-lite-preview-09-2025",
        "gemini-2.5-flash-lite-preview-09-2025",
        "models/gemini-2.5-flash-lite-preview-06-17",
        "gemini-2.5-flash-lite-preview-06-17",
    ],
    "Gemini 2.0 Flash": [
        "models/gemini-2.0-flash",
        "gemini-2.0-flash",
        "models/gemini-2.0-flash-001",
        "gemini-2.0-flash-001",
        "models/gemini-2.0-flash-exp",
        "gemini-2.0-flash-exp",
    ],
    "Gemini 2.0 Pro (Experimental)": [
        "models/gemini-2.0-pro-exp",
        "gemini-2.0-pro-exp",
        "models/gemini-2.0-pro-exp-02-05",
        "gemini-2.0-pro-exp-02-05",
    ],
    "Gemini Flash (Latest Alias)": [
        "models/gemini-flash-latest",
        "gemini-flash-latest",
    ],
    "Gemini Pro (Latest Alias)": [
        "models/gemini-pro-latest",
        "gemini-pro-latest",
    ],
}
DEFAULT_MODEL_LABEL = "Gemini 2.5 Pro"


def safe_rerun() -> None:
    """Trigger a Streamlit rerun, compatible with older versions."""
    try:
        st.rerun()
    except AttributeError:  # pragma: no cover - older Streamlit versions
        st.experimental_rerun()


def _normalize_credential(value: Optional[str]) -> Optional[str]:
    if isinstance(value, str):
        stripped = value.strip()
        if stripped:
            return stripped
    return None


def get_secret_auth_credentials() -> Tuple[Optional[str], Optional[str]]:
    try:
        secrets_obj = st.secrets
    except StreamlitSecretNotFoundError:
        return None, None
    except Exception:  # pragma: no cover - defensive
        return None, None

    auth_section: Optional[Dict[str, Any]] = None
    if isinstance(secrets_obj, dict):
        auth_section = secrets_obj.get("auth")
    else:
        # Streamlit's Secrets object behaves like a mapping.
        auth_section = getattr(secrets_obj, "get", lambda _key, _default=None: None)("auth")

    if not isinstance(auth_section, dict):
        return None, None

    username = auth_section.get("username") or auth_section.get("id")
    password = auth_section.get("password") or auth_section.get("pass")
    return _normalize_credential(str(username)) if username is not None else None, _normalize_credential(
        str(password)
    ) if password is not None else None


def get_configured_auth_credentials() -> Tuple[Optional[str], Optional[str]]:
    """Return the currently configured Basic auth credentials."""
    session_username = _normalize_credential(st.session_state.get("auth_username"))
    session_password = _normalize_credential(st.session_state.get("auth_password"))
    if session_username and session_password:
        return session_username, session_password
    return get_secret_auth_credentials()


def render_basic_auth_settings(
    form_prefix: str,
    *,
    caption: Optional[str] = "Basic認証のID・パスワードを設定してください。空欄の場合は未設定として扱われます。",
    expand_when_missing: bool = True,
) -> None:
    current_username, current_password = get_configured_auth_credentials()
    expanded = expand_when_missing and not (current_username and current_password)

    with st.expander("Basic認証の設定", expanded=expanded):
        if caption:
            st.caption(caption)
        with st.form(f"{form_prefix}_auth_form", clear_on_submit=False):
            username_input = st.text_input(
                "Basic 認証 ID",
                value=current_username or "",
            )
            password_input = st.text_input(
                "Basic 認証 パスワード",
                value=current_password or "",
                type="password",
            )
            submit_col, clear_col = st.columns(2)
            submitted = submit_col.form_submit_button("設定を保存")
            cleared = clear_col.form_submit_button("クリア")

        if submitted:
            normalized_username = _normalize_credential(username_input)
            normalized_password = _normalize_credential(password_input)
            st.session_state["auth_username"] = normalized_username
            st.session_state["auth_password"] = normalized_password
            st.session_state["authenticated"] = False
            st.success("設定を保存しました。")
            safe_rerun()
        elif cleared:
            st.session_state["auth_username"] = None
            st.session_state["auth_password"] = None
            st.session_state["authenticated"] = False
            st.info("設定をクリアしました。")
            safe_rerun()


def require_login() -> None:
    """Render a simple login form and block the rest of the app until authenticated."""
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if st.session_state["authenticated"]:
        return

    st.title("ログイン")

    username, password = get_configured_auth_credentials()
    if not username or not password:
        st.info("ログイン情報が未設定です。管理者に連絡してください。")
        st.stop()
        return

    with st.form("login_form", clear_on_submit=False):
        input_username = st.text_input("ID")
        input_password = st.text_input("PASS", type="password")
        submitted = st.form_submit_button("ログイン")

    if submitted:
        if input_username == username and input_password == password:
            st.session_state["authenticated"] = True
            st.success("ログインしました。")
            safe_rerun()
            return
        st.error("IDまたはPASSが正しくありません。")
    st.stop()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    if PdfReader is None:
        return "(PDFのテキスト抽出に必要なライブラリがインストールされていません: pypdf)"
    reader = PdfReader(BytesIO(file_bytes))
    chunks: List[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception as exc:  # pragma: no cover - defensive
            chunks.append(f"[PDFページの抽出に失敗しました: {exc}]")
    return "\n".join(chunks).strip()


def extract_images_from_pdf(file_bytes: bytes, source_name: str) -> List[Dict[str, Any]]:
    images: List[Dict[str, Any]] = []
    if PdfReader is None:
        return images

    reader = PdfReader(BytesIO(file_bytes))
    source_base = os.path.splitext(source_name)[0]
    for page_number, page in enumerate(reader.pages, start=1):
        page_images = getattr(page, "images", []) or []
        for image_index, image in enumerate(page_images, start=1):
            data = getattr(image, "data", None)
            if not data:
                continue
            ext = (getattr(image, "ext", None) or getattr(image, "name", "")).split(".")[-1].lower()
            if not ext or len(ext) > 4:
                ext = "png"
            image_id = f"img_{uuid.uuid4().hex}"
            filename = f"{source_base}_p{page_number}_img{image_index}.{ext}"
            images.append(
                {
                    "id": image_id,
                    "bytes": data,
                    "filename": filename,
                    "ext": ext,
                    "source": source_name,
                    "page": page_number,
                }
            )
    return images


def extract_text_from_docx(file_bytes: bytes) -> str:
    if Document is None:
        return "(DOCXのテキスト抽出に必要なライブラリがインストールされていません: python-docx)"
    document = Document(BytesIO(file_bytes))
    return "\n".join(p.text for p in document.paragraphs).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("utf-8", errors="ignore")


def extract_text(upload) -> Dict[str, Any]:
    file_bytes = upload.getvalue()
    name = upload.name
    suffix = name.split(".")[-1].lower() if "." in name else ""

    if suffix == "pdf":
        content = extract_text_from_pdf(file_bytes)
        images = extract_images_from_pdf(file_bytes, name)
    elif suffix in IMAGE_SUFFIXES:
        content = "(画像ファイルです。内容は自動解析されていません。必要に応じて指示欄で補足してください。)"
        images = [
            {
                "id": f"img_{uuid.uuid4().hex}",
                "bytes": file_bytes,
                "filename": name,
                "ext": suffix,
                "source": name,
                "page": None,
            }
        ]
    elif suffix in {"docx", "doc"}:
        if suffix == "doc":
            content = "(旧形式のDOCファイルには未対応です。DOCX形式に変換してください。)"
            images = []
        else:
            content = extract_text_from_docx(file_bytes)
            images = []
    elif suffix in {"txt", "md"}:
        content = extract_text_from_txt(file_bytes)
        images = []
    else:
        content = "(未対応のファイル形式です。テキストとして解釈できませんでした。)"
        images = []

    return {"name": name, "content": content, "images": images}


def build_prompt(base_prompt: str, instructions: str, sources: List[Dict[str, Any]]) -> str:
    prompt_lines = [base_prompt.strip(), "", "# 追加指示"]
    prompt_lines.append(instructions.strip() if instructions.strip() else "特別な追加指示はありません。")
    prompt_lines.append("")

    if sources:
        prompt_lines.append("# 参考資料")
        for source in sources:
            prompt_lines.append(f"## {source['name']}")
            prompt_lines.append(source["content"].strip() or "(本文なし)")
            if source.get("images"):
                prompt_lines.append("### 添付画像リスト")
                for image in source["images"]:
                    desc = f"- {image['filename']}"
                    if image.get("page"):
                        desc += f" (ページ {image['page']})"
                    prompt_lines.append(desc)
            prompt_lines.append("")
    else:
        prompt_lines.append("(参考資料はアップロードされていません。)")

    prompt_lines.append("これらの情報をもとに、プレスリリース記事を作成してください。")
    prompt_lines.append("")
    prompt_lines.append(OUTPUT_FORMAT_INSTRUCTIONS)
    return "\n".join(prompt_lines).strip()


def create_docx_bytes(
    headline: str,
    subheadline: str,
    article: str,
    images: List[Dict[str, Any]],
) -> bytes:
    if Document is None:
        raise RuntimeError("DOCX出力には python-docx が必要です。")

    doc = Document()
    doc.add_heading(headline, level=0)
    doc.add_heading(subheadline, level=1)

    normalized = article.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [block.strip() for block in normalized.split("\n\n")]
    positions: Dict[int, List[Dict[str, Any]]] = {}
    for item in images:
        pos = int(item.get("position", 0))
        pos = max(0, min(pos, len(paragraphs)))
        positions.setdefault(pos, []).append(item)

    def add_image_to_doc(image_item: Dict[str, Any]) -> None:
        image = image_item["image"]
        caption = image_item.get("caption")
        raw_bytes = image["bytes"]

        def insert_picture(stream: BytesIO) -> None:
            stream.seek(0)
            if Inches is not None:
                doc.add_picture(stream, width=Inches(5.5))
            else:
                doc.add_picture(stream)

        try:
            insert_picture(BytesIO(raw_bytes))
        except Exception as original_exc:
            if Image is None:
                raise RuntimeError(
                    f"画像 '{image.get('filename')}' をDOCXに追加できませんでした。Pillow のインストールを検討してください。"
                ) from original_exc
            try:
                with Image.open(BytesIO(raw_bytes)) as pil_image:
                    converted = BytesIO()
                    # JPEG で保存すると比較的軽量に保てる
                    save_kwargs = {"format": "PNG" if pil_image.mode in ("RGBA", "LA") else "JPEG"}
                    if save_kwargs["format"] == "JPEG" and pil_image.mode not in ("RGB", "L"):
                        pil_image = pil_image.convert("RGB")
                    pil_image.save(converted, **save_kwargs)
                    insert_picture(converted)
            except Exception as pil_exc:  # pragma: no cover - sample dependent
                raise RuntimeError(
                    f"画像 '{image.get('filename')}' の変換に失敗しました。"
                ) from pil_exc
        if caption:
            try:
                doc.add_paragraph(caption, style="Caption")
            except (KeyError, ValueError):
                doc.add_paragraph(caption)

    for image_item in positions.get(0, []):
        add_image_to_doc(image_item)

    for index, block in enumerate(paragraphs, start=1):
        cleaned_block = block.replace("  \n", " ").replace("\t", " ")
        if cleaned_block:
            para = doc.add_paragraph(cleaned_block)
            for inline in para.runs:
                inline.font.name = None
                inline.font.size = None
        else:
            doc.add_paragraph("")
        for image_item in positions.get(index, []):
            add_image_to_doc(image_item)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def configure_genai(api_key: str):
    if genai is None:
        raise RuntimeError(
            "google-generativeai がインポートできませんでした。 'pip install google-generativeai' を実行してください。"
        ) from _import_error
    try:
        genai.configure(api_key=api_key, api_version="v1")
    except TypeError:
        # 旧バージョンの SDK は api_version を受け付けない。
        genai.configure(api_key=api_key)


def _candidate_model_names(model_name: str) -> List[str]:
    candidates = [model_name]
    if model_name.startswith("models/"):
        candidates.append(model_name.split("models/", 1)[1])
    else:
        candidates.append(f"models/{model_name}")
    # preserve order without duplicates
    seen = set()
    ordered: List[str] = []
    for name in candidates:
        if name and name not in seen:
            ordered.append(name)
            seen.add(name)
    return ordered


def _extract_text_from_response(response) -> str:
    if hasattr(response, "text") and response.text:
        return response.text
    if getattr(response, "candidates", None):
        texts: List[str] = []
        for candidate in response.candidates:
            content = getattr(candidate, "content", None)
            if not content:
                continue
            for part in getattr(content, "parts", []) or []:
                text = getattr(part, "text", "")
                if text:
                    texts.append(text)
        if texts:
            return "\n".join(texts)
    return ""


def _strip_code_fences(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```") and cleaned.endswith("```"):
        inner = cleaned[3:-3]
        if "\n" in inner:
            first_line, rest = inner.split("\n", 1)
            if first_line.strip().lower() in {"json", ""}:
                return rest.strip()
            return inner.strip()
        return inner.strip()
    return cleaned


def parse_model_output(raw_text: str) -> Dict[str, Any]:
    cleaned = _strip_code_fences(raw_text)
    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise ValueError("Gemini からの出力を JSON として解析できませんでした。") from exc

    def _to_plain_text(value: str) -> str:
        text = value.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"</(p|div|section|article|figure)>", "\n\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<(p|div|section|article|figure)[^>]*>", "", text, flags=re.IGNORECASE)
        text = re.sub(r"<h[1-6][^>]*>", "\n\n", text, flags=re.IGNORECASE)
        text = re.sub(r"</h[1-6]>", "\n\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<figcaption[^>]*>", "", text, flags=re.IGNORECASE)
        text = re.sub(r"</figcaption>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<li[^>]*>", "- ", text, flags=re.IGNORECASE)
        text = re.sub(r"</li>", "\n", text, flags=re.IGNORECASE)
        text = HTML_TAG_RE.sub("", text)
        text = re.sub(r"^\s*#+\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = text.replace("• ", "").replace("- ", "")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    headlines = data.get("headlines")
    subheadlines = data.get("subheadlines")
    article = data.get("article")

    if not isinstance(headlines, list) or len(headlines) != 10 or not all(isinstance(item, str) for item in headlines):
        raise ValueError("見出し候補が10件の文字列配列として得られませんでした。")
    if not isinstance(subheadlines, list) or len(subheadlines) != 10 or not all(isinstance(item, str) for item in subheadlines):
        raise ValueError("小見出し候補が10件の文字列配列として得られませんでした。")
    if not isinstance(article, str) or not article.strip():
        raise ValueError("本文が空、または文字列として取得できませんでした。")

    article_plain = _to_plain_text(article)

    return {
        "headlines": [_to_plain_text(item) for item in headlines],
        "subheadlines": [_to_plain_text(item) for item in subheadlines],
        "article": article_plain,
        "article_raw": article,
    }


def generate_press_release(model_names: Iterable[str], prompt: str) -> str:
    last_error: Optional[Exception] = None
    for model_name in model_names:
        for candidate_name in _candidate_model_names(model_name):
            try:
                model = genai.GenerativeModel(candidate_name)
                response = model.generate_content(prompt)
                text = _extract_text_from_response(response)
                if text:
                    return text
                return "生成結果をテキストとして取得できませんでした。"
            except NotFound as exc:
                last_error = exc
                continue
    if last_error:
        raise last_error
    raise RuntimeError("Gemini モデルの呼び出しに失敗しました。")


def main():
    st.set_page_config(page_title="めざましメディア・プレスリリース生成", page_icon="📰", layout="wide")
    require_login()
    st.title("めざましメディア プレスリリース生成ツール")
    st.write(
        "アップロードした資料と指示をもとに、Gemini がプレスリリース案を作成します。"
    )

    if "generation" not in st.session_state:
        st.session_state["generation"] = None
    if "selection_finalized" not in st.session_state:
        st.session_state["selection_finalized"] = False
    if "selected_headline" not in st.session_state:
        st.session_state["selected_headline"] = ""
    if "selected_subheadline" not in st.session_state:
        st.session_state["selected_subheadline"] = ""
    if "available_images" not in st.session_state:
        st.session_state["available_images"] = []
    if "selected_images_config" not in st.session_state:
        st.session_state["selected_images_config"] = []

    with st.sidebar:
        st.header("Gemini 設定")
        if "gemini_api_key" not in st.session_state:
            initial_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY") or ""
            if initial_key:
                st.session_state["gemini_api_key"] = initial_key.strip()

        query_params = st.query_params
        is_admin_mode = query_params.get("admin") == "1"

        if is_admin_mode:
            render_basic_auth_settings(
                "sidebar",
                caption="Basic認証のID・パスワードを更新できます。変更すると再ログインが必要です。",
                expand_when_missing=False,
            )

        api_key = st.session_state.get("gemini_api_key", "")
        if is_admin_mode:
            api_key_input = st.text_input(
                "APIキー (管理者専用)",
                value="",
                type="password",
                placeholder="Gemini の APIキーを入力",
                label_visibility="collapsed",
            )
            if api_key_input:
                st.session_state["gemini_api_key"] = api_key_input.strip()
            if st.button("APIキーをクリア", key="clear_api_key"):
                st.session_state["gemini_api_key"] = ""
            api_key = st.session_state.get("gemini_api_key", "")
        elif not api_key:
            st.info("APIキーは管理者が設定します。必要な場合は管理者に連絡してください。")
        model_labels = list(MODEL_OPTIONS.keys())
        default_index = model_labels.index(DEFAULT_MODEL_LABEL) if DEFAULT_MODEL_LABEL in model_labels else 0
        selected_model_label = st.selectbox("モデル", model_labels, index=default_index)
        model_candidates = MODEL_OPTIONS[selected_model_label]
        with st.expander("プロンプトを編集する", expanded=False):
            base_prompt = st.text_area("ベースプロンプト", value=BASE_PROMPT, height=240)

    uploads = st.file_uploader(
        "資料ファイルをアップロード (PDF, TXT, DOCX, 画像など)",
        type=[
            "pdf",
            "txt",
            "md",
            "docx",
            "doc",
            "png",
            "jpg",
            "jpeg",
            "webp",
            "bmp",
            "gif",
        ],
        accept_multiple_files=True,
    )

    instructions = st.text_area(
        "記事の方向性や盛り込みたいポイント（自由入力）",
        help="記事のトーンや必須要素など、Gemini に伝えたい内容を入力してください。",
        height=200,
    )

    extracted_sources: List[Dict[str, Any]] = []
    if uploads:
        st.subheader("抽出されたテキスト")
        for upload in uploads:
            source = extract_text(upload)
            extracted_sources.append(source)
            with st.expander(source["name"], expanded=False):
                st.write(source["content"] or "(テキストを抽出できませんでした)")
                if source.get("images"):
                    st.markdown("**抽出された画像プレビュー**")
                    for img in source["images"]:
                        st.image(
                            img["bytes"],
                            caption=f"{source['name']} / ページ {img['page']}",
                            width=300,
                        )

    disabled = not api_key or (uploads is None or len(uploads) == 0)
    if st.button("Gemini に送信して記事を生成", disabled=disabled):
        if not api_key:
            st.error("Gemini API Key を入力してください。")
            return
        if not extracted_sources:
            st.error("少なくとも1つの資料をアップロードしてください。")
            return

        for key in [
            key for key in st.session_state.keys() if key.startswith("img_include_") or key.startswith("img_pos_") or key.startswith("img_caption_")
        ]:
            del st.session_state[key]

        prompt = build_prompt(base_prompt, instructions, extracted_sources)
        try:
            with st.spinner("Gemini が記事を作成しています..."):
                configure_genai(api_key)
                raw_result = generate_press_release(model_candidates, prompt)
                parsed_result = parse_model_output(raw_result)
        except Exception as exc:  # pragma: no cover - interactive feedback
            st.error(f"生成中にエラーが発生しました: {exc}")
            return

        st.session_state["generation"] = parsed_result
        st.session_state["selection_finalized"] = False
        st.session_state["selected_headline"] = ""
        st.session_state["selected_subheadline"] = ""
        st.session_state["headline_choice"] = parsed_result["headlines"][0]
        st.session_state["subheadline_choice"] = parsed_result["subheadlines"][0]
        st.session_state["available_images"] = [
            image
            for source in extracted_sources
            for image in source.get("images", [])
        ]
        st.session_state["selected_images_config"] = []
        st.success("候補を生成しました。見出しと小見出しを選択してください。")

    generation = st.session_state.get("generation")
    if generation:
        st.subheader("見出しと小見出しの候補")
        if st.session_state.get("selection_finalized"):
            final_headline = st.session_state.get("selected_headline", "").strip()
            final_subheadline = st.session_state.get("selected_subheadline", "").strip()
            if not final_headline or not final_subheadline:
                st.warning("見出しまたは小見出しの選択が完了していません。")
            else:
                spinner_context = (
                    st.spinner("見出しと小見出しを確定しています...")
                    if st.session_state.pop("finalizing_indicator", False)
                    else nullcontext()
                )
                with spinner_context:
                    st.success("見出しと小見出しを確定しました。")
                    st.markdown(f"### {final_headline}")
                    st.markdown(f"#### {final_subheadline}")

                    paragraphs = generation["article"].split("\n\n")
                    available_image_map = {
                        image["id"]: image for image in st.session_state.get("available_images", [])
                    }
                    resolved_images = []
                    for config in st.session_state.get("selected_images_config", []):
                        image = available_image_map.get(config["image_id"])
                        if image:
                            resolved_images.append({"image": image, **config})

                    position_map: Dict[int, List[Dict[str, Any]]] = {}
                    for item in resolved_images:
                        pos = int(item.get("position", 0))
                        pos = max(0, min(pos, len(paragraphs)))
                        position_map.setdefault(pos, []).append(item)

                    def render_image(item: Dict[str, Any]):
                        image = item["image"]
                        caption = item.get("caption") or image.get("filename")
                        st.image(
                            image["bytes"],
                            caption=f"{caption} (出典: {image['source']})" if caption else f"出典: {image['source']}",
                            width=300,
                        )

                    if position_map.get(0):
                        for img_item in position_map[0]:
                            render_image(img_item)

                    for idx, paragraph in enumerate(paragraphs, start=1):
                        st.markdown(paragraph.replace("\n", "  \n"))
                        for img_item in position_map.get(idx, []):
                            render_image(img_item)

                    text_parts = [final_headline, final_subheadline]
                    if position_map.get(0):
                        for item in position_map[0]:
                            caption = item.get("caption") or item["image"].get("filename")
                            text_parts.append(f"[画像: {caption}]")
                    for idx, paragraph in enumerate(paragraphs, start=1):
                        text_parts.append(paragraph)
                        for item in position_map.get(idx, []):
                            caption = item.get("caption") or item["image"].get("filename")
                            text_parts.append(f"[画像: {caption}]")
                    final_text = "\n\n".join(part for part in text_parts if part)
                    st.download_button(
                        "テキストをダウンロード",
                        data=final_text,
                        file_name="press_release_draft.txt",
                        mime="text/plain",
                        key="download_final",
                    )
                    docx_bytes = None
                    docx_error = None
                    if Document is None:
                        docx_error = "DOCX形式でのダウンロードには python-docx のインストールが必要です。"
                    else:
                        try:
                            docx_bytes = create_docx_bytes(
                                final_headline,
                                final_subheadline,
                                generation["article"],
                                resolved_images,
                            )
                        except Exception as exc:
                            docx_error = f"DOCXファイルの生成に失敗しました: {exc}"

                    if docx_bytes:
                        st.download_button(
                            "DOCX をダウンロード",
                            data=docx_bytes,
                            file_name="press_release_draft.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx",
                        )
                    elif docx_error:
                        st.info(docx_error)
                if st.button("見出しの選択をやり直す", key="reset_selection"):
                    st.session_state["selection_finalized"] = False
        else:
            st.session_state.setdefault("headline_choice", generation["headlines"][0])
            st.session_state.setdefault("subheadline_choice", generation["subheadlines"][0])
            headline_choice = st.selectbox(
                "見出し候補 (10件)",
                generation["headlines"],
                key="headline_choice",
            )
            subheadline_choice = st.selectbox(
                "小見出し候補 (10件)",
                generation["subheadlines"],
                key="subheadline_choice",
            )

            st.markdown("#### 生成された記事案")
            st.write(generation["article"])

            available_images = st.session_state.get("available_images", [])
            if available_images:
                st.markdown("#### 挿入する画像の選択")
                paragraphs = generation["article"].split("\n\n")
                insertion_points = [(0, "見出し直後")]
                for idx, paragraph in enumerate(paragraphs, start=1):
                    preview = paragraph.replace("\n", " ")[:40]
                    label = f"本文 第{idx}段落の後"
                    if preview:
                        label += f" ：{preview}"
                    insertion_points.append((idx, label))

                for image in available_images:
                    image_id = image["id"]
                    include_key = f"img_include_{image_id}"
                    default_include = st.session_state.get(include_key, False)
                    with st.container():
                        st.image(
                            image["bytes"],
                            caption=f"{image['source']} / ページ {image['page']}"
                            if image.get("page")
                            else image["filename"],
                            width=300,
                        )
                        include = st.checkbox(
                            "この画像を挿入する",
                            value=default_include,
                            key=include_key,
                        )
                        if include:
                            position_options = [choice[0] for choice in insertion_points]
                            current_position = st.session_state.get(f"img_pos_{image_id}", position_options[0])

                            position = st.selectbox(
                                "挿入位置",
                                options=position_options,
                                format_func=lambda value, choices=insertion_points: next(
                                    label for val, label in choices if val == value
                                ),
                                key=f"img_pos_{image_id}",
                                index=position_options.index(current_position)
                                if current_position in position_options
                                else 0,
                            )
                            caption_value = st.text_input(
                                "キャプション (任意)",
                                key=f"img_caption_{image_id}",
                                value=st.session_state.get(f"img_caption_{image_id}", ""),
                                help="本文中に表示する補足説明があれば入力してください。",
                            )
                        else:
                            for suffix in ("pos", "caption"):
                                st.session_state.pop(f"img_{suffix}_{image_id}", None)

            st.write("候補は各10件です。選択後に確定ボタンを押してください。")
            if st.button("見出しと小見出しを確定する", key="confirm_selection"):
                with st.spinner("見出しと小見出しを確定しています..."):
                    st.session_state["selected_headline"] = headline_choice
                    st.session_state["selected_subheadline"] = subheadline_choice
                    available_image_map = {
                        image["id"]: image for image in st.session_state.get("available_images", [])
                    }
                    selected_images_config: List[Dict[str, Any]] = []
                    for image_id in available_image_map:
                        if st.session_state.get(f"img_include_{image_id}"):
                            position = int(st.session_state.get(f"img_pos_{image_id}", 0))
                            caption_value = st.session_state.get(f"img_caption_{image_id}", "").strip()
                            selected_images_config.append(
                                {
                                    "image_id": image_id,
                                    "position": position,
                                    "caption": caption_value,
                                }
                            )
                    st.session_state["selected_images_config"] = selected_images_config
                    st.session_state["selection_finalized"] = True
                    st.session_state["finalizing_indicator"] = True
                st.rerun()


if __name__ == "__main__":
    main()
